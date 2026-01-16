"""
文档文本提取工具模块

支持PDF、Word、TXT、Markdown等格式的文本提取。
"""

import os
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO


def extract_text_from_file(file_path: str) -> str | None:
    """根据文件类型提取文本（本地文件路径）"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_from_word(file_path)
    elif ext in [".txt", ".md"]:
        return extract_from_text(file_path)
    else:
        return None


def extract_from_pdf(file_path: str) -> str | None:
    """从PDF提取文本（本地文件）"""
    try:
        text_parts = []
        doc = fitz.open(file_path)
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "".join(map(str, text_parts)).strip()
    except Exception as e:
        print(f"PDF解析失败: {e}")
        return None


def extract_from_word(file_path: str) -> str | None:
    """从Word文档提取文本（本地文件）"""
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += "\n" + row_text
        return text.strip()
    except Exception as e:
        print(f"Word解析失败: {e}")
        return None


def extract_from_text(file_path: str) -> str | None:
    """从纯文本提取内容（本地文件）"""
    try:
        encodings = ["utf-8", "gbk", "gb2312", "latin1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read().strip()
            except UnicodeDecodeError:
                continue
        return None
    except Exception as e:
        print(f"文本读取失败: {e}")
        return None


def extract_from_upload(file_storage) -> str | None:
    """从上传的文件对象提取文本（支持内存文件）"""
    filename = file_storage.filename or ""
    ext = os.path.splitext(filename)[1].lower()
    if not ext:
        return None

    try:
        if ext == ".pdf":
            return extract_from_pdf_memory(file_storage)
        elif ext in [".docx", ".doc"]:
            return extract_from_word_memory(file_storage)
        elif ext in [".txt", ".md"]:
            return file_storage.read().decode("utf-8").strip()
        else:
            return None
    except Exception as e:
        print(f"文件解析失败: {e}")
        return None


def extract_from_pdf_memory(file_storage) -> str | None:
    """从内存PDF提取文本"""
    try:
        # 保存原始位置
        original_position = file_storage.tell()
        file_storage.seek(0)
        content = file_storage.read()
        file_storage.seek(original_position)  # 恢复原始位置

        text_parts = []
        doc = fitz.open(stream=content, filetype="pdf")
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "".join(map(str, text_parts)).strip()
    except Exception as e:
        print(f"PDF解析失败: {e}")
        # 尝试恢复指针
        try:
            file_storage.seek(0)
        except:
            pass
        return None


def extract_from_word_memory(file_storage) -> str | None:
    """从内存Word文档提取文本"""
    try:
        original_position = file_storage.tell()
        file_storage.seek(0)
        doc = Document(file_storage)
        text = "\n".join([para.text for para in doc.paragraphs])

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += "\n" + row_text

        file_storage.seek(original_position)  # 恢复原始位置
        return text.strip()
    except Exception as e:
        print(f"Word解析失败: {e}")
        try:
            file_storage.seek(0)
        except:
            pass
        return None


def extract_text_preview(text: str, max_length: int = 500) -> str:
    """提取文本预览（前N个字符）"""
    if not text:
        return ""
    return text[:max_length] + ("..." if len(text) > max_length else "")


def count_text_words(text: str) -> int:
    """统计文本字数"""
    if not text:
        return 0
    return len(text)


def is_scanned_pdf(file_storage) -> bool:
    """判断PDF是否为扫描件（通过检测是否有可提取文本）"""
    try:
        file_storage.seek(0)
        content = file_storage.read()
        file_storage.seek(0)

        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()

        text = "".join(map(str, text_parts))
        # 如果提取的文本很少，可能是扫描件
        return len(text.strip()) < 50
    except Exception:
        return True


if __name__ == "__main__":
    # 测试
    import sys

    if len(sys.argv) > 1:
        result = extract_text_from_file(sys.argv[1])
        if result:
            print(f"提取成功，字数: {len(result)}")
            print(f"预览: {extract_text_preview(result, 200)}")
        else:
            print("提取失败")
